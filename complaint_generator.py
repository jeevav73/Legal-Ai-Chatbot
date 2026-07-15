"""
complaint_generator.py
Takes structured intake details from the user, asks the local LLM (Aalap) to
draft formal complaint text, and renders it into a .docx file using python-docx.
"""

from datetime import date

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config
from rag_engine import generate_text, retrieve_context


COMPLAINT_DRAFT_INSTRUCTION = """You are drafting a formal written police/legal complaint
on behalf of a citizen in India. Use a formal, respectful, factual tone. Do not invent
facts beyond what is given. Structure it as: Addressee line, Subject line, salutation,
numbered factual paragraphs, and a closing paragraph requesting specific action, followed
by "Yours faithfully," and a signature line. Do not include any case law citations unless
given. Keep it under 400 words."""


def draft_complaint_text(details: dict) -> str:
    """Calls the local LLM to turn structured intake details into formal complaint prose."""
    context_chunks = retrieve_context(
        f"How to draft a complaint about: {details.get('incident_summary', '')}"
    )
    context_text = "\n\n".join(chunk for chunk, _ in context_chunks)

    prompt = (
        f"{COMPLAINT_DRAFT_INSTRUCTION}\n\n"
        f"RELEVANT LEGAL CONTEXT:\n{context_text}\n\n"
        f"COMPLAINANT DETAILS:\n"
        f"Name: {details.get('name')}\n"
        f"Address: {details.get('address')}\n"
        f"Contact: {details.get('contact')}\n"
        f"Addressed to: {details.get('addressee')}\n"
        f"Date of incident: {details.get('incident_date')}\n"
        f"Location: {details.get('location')}\n"
        f"Incident summary: {details.get('incident_summary')}\n"
        f"Accused/opposite party (if known): {details.get('accused')}\n"
        f"Witnesses: {details.get('witnesses')}\n"
        f"Evidence available: {details.get('evidence')}\n"
        f"Relief requested: {details.get('relief_requested')}\n\n"
        f"Now write the full formal complaint text:"
    )

    return generate_text(prompt, max_new_tokens=600)


def build_docx(details: dict, complaint_text: str) -> str:
    """Renders the drafted complaint into a formatted .docx file and returns its path."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Header block: complainant's own details, right-aligned like a letterhead
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run(f"{details.get('name')}\n{details.get('address')}\n{details.get('contact')}\n"
                    f"Date: {date.today().strftime('%d-%m-%Y')}").font.size = Pt(10)

    doc.add_paragraph()  # spacer

    doc.add_paragraph(f"To,\n{details.get('addressee')}")
    doc.add_paragraph()

    subject = doc.add_paragraph()
    subject.add_run(f"Subject: Complaint regarding {details.get('incident_summary', '')[:80]}").bold = True

    doc.add_paragraph()
    doc.add_paragraph("Respected Sir/Madam,")
    doc.add_paragraph()

    for para in complaint_text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.add_paragraph()
    doc.add_paragraph("Yours faithfully,")
    doc.add_paragraph(details.get("name"))

    # Safety footer — always included
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "NOTE: This is an AI-assisted draft generated for reference. Please have it "
        "reviewed by a licensed advocate before filing."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)

    filename_safe_name = (details.get("name") or "complaint").replace(" ", "_")
    output_path = config.GENERATED_COMPLAINTS_DIR / f"complaint_{filename_safe_name}_{date.today().isoformat()}.docx"
    doc.save(str(output_path))
    return str(output_path)


def generate_complaint(details: dict) -> str:
    """End-to-end: draft text with the LLM, then render into a .docx. Returns file path."""
    complaint_text = draft_complaint_text(details)
    return build_docx(details, complaint_text)


if __name__ == "__main__":
    # Quick manual test
    sample_details = {
        "name": "Ravi Kumar",
        "address": "12, Gandhi Street, Chennai",
        "contact": "9876543210",
        "addressee": "The Superintendent of Police, Chennai District",
        "incident_date": "05-07-2026",
        "location": "Attur Police Station",
        "incident_summary": "A false case was registered against me based on a personal "
                             "dispute with my neighbour, without any investigation.",
        "accused": "N/A",
        "witnesses": "Mr. Suresh (neighbour), Ms. Latha (shopkeeper)",
        "evidence": "WhatsApp messages, 2 witness statements",
        "relief_requested": "Kindly look into the matter and ensure a fair, unbiased "
                             "investigation into the FIR registered against me.",
    }
    path = generate_complaint(sample_details)
    print(f"Complaint saved to: {path}")
